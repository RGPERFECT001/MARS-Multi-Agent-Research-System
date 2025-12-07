"""Word Agent for converting reports to Word documents."""

import logging
import os
from typing import Dict, Any

# Ensure python-docx is installed: pip install python-docx
try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    raise ImportError("Please install python-docx: pip install python-docx")

# Adjust import based on your repo structure (models.py is in root)
from models import AgentState

logger = logging.getLogger(__name__)

class WordAgent:
    """Agent responsible for converting the final report into a Word document."""
    
    def __init__(self):
        """Initialize the Word Agent."""
        self.name = "Word Agent"
        self.output_dir = "outputs"
        
        # Ensure output directory exists
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
            
        logger.info(f"Initialized {self.name}")
    
    def convert_to_word(self, state) -> Dict[str, Any]:
        """
        Convert the final report to a Word document.
        Args:
            state: Current agent state (object or dict)
        """
        logger.info(f"{self.name} converting report to Word document")
        
        # Flexible state handling (dict or object)
        if isinstance(state, dict):
            report_content = state.get('final_report') or state.get('draft_report')
            topic = state.get('user_topic', 'Research Report')
        else:
            report_content = getattr(state, 'final_report', None) or getattr(state, 'draft_report', None)
            topic = getattr(state, 'user_topic', 'Research Report')
        
        if not report_content:
            logger.warning("No report content to convert")
            return {}
            
        try:
            doc = Document()
            doc.add_heading(topic, 0)
            
            lines = report_content.split('\n')
            for line in lines:
                line = line.strip()
                if not line: continue
                    
                # Simple Markdown Parsing
                if line.startswith('# '): doc.add_heading(line[2:], 1)
                elif line.startswith('## '): doc.add_heading(line[3:], 2)
                elif line.startswith('### '): doc.add_heading(line[4:], 3)
                elif line.startswith('#### '): doc.add_heading(line[5:], 4)
                elif line.startswith('- ') or line.startswith('* '):
                    p = doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith('1. '):
                    p = doc.add_paragraph(line[3:], style='List Number')
                else:
                    doc.add_paragraph(line)
            
            # Sanitize filename
            safe_topic = "".join([c for c in topic if c.isalpha() or c.isdigit() or c in (' ', '-', '_')]).rstrip()
            filename = f"{safe_topic.replace(' ', '_')}_report.docx"
            filepath = os.path.join(self.output_dir, filename)
            
            doc.save(filepath)
            logger.info(f"Report saved to {filepath}")
            
            return {"word_document_path": filepath, "filename": filename}
            
        except Exception as e:
            logger.error(f"{self.name} failed to convert report: {e}")
            return {}